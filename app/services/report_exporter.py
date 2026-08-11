"""
report_exporter.py — Generates PDF, DOCX, and TXT compliance reports.
PDF generation uses ReportLab (industry-standard).
DOCX generation uses python-docx.
All generation runs in background threads — this module has no UI dependencies.
"""
from __future__ import annotations
import os
import datetime
from typing import List

from app.validators.base import Finding
from app.services.scanner import ScanResult
from app.utils.path_helper import get_reports_dir


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def _timestamp() -> str:
    return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")


def _human_timestamp() -> str:
    return datetime.datetime.now().strftime("%d %B %Y, %H:%M:%S")


def _severity_color_rgb(severity: str):
    """Return (R, G, B) floats 0-1 for a severity level."""
    return {
        "Critical":    (0.776, 0.157, 0.157),
        "Warning":     (0.937, 0.424, 0.024),
        "Information": (0.012, 0.533, 0.820),
    }.get(severity, (0.4, 0.4, 0.4))


def _grade_color(grade: str):
    return {
        "Compliant":           (0.18, 0.49, 0.20),
        "Partially Compliant": (0.937, 0.424, 0.024),
        "Non-Compliant":       (0.776, 0.157, 0.157),
    }.get(grade, (0.4, 0.4, 0.4))


# ─────────────────────────────────────────────────────────────────────────────
# TXT Export
# ─────────────────────────────────────────────────────────────────────────────
def export_txt(result: ScanResult, filename_base: str = None, filepath: str = None) -> str:
    """Generate a TXT report. Returns the saved file path."""
    if not filepath:
        reports_dir = get_reports_dir()
        filename_base = filename_base or "compliance_report"
        filepath = os.path.join(reports_dir, f"{filename_base}_{_timestamp()}.txt")

    if result.grade == "Compliant":
        summary_text = (
            "The scanned document achieved a compliance score of "
            f"{result.score:.1f}% and is rated Compliant. "
            "No critical issues were identified. Minor improvements may still apply."
        )
    elif result.grade == "Partially Compliant":
        summary_text = (
            f"The scanned document scored {result.score:.1f}% and is rated "
            "Partially Compliant. "
            f"There are {result.critical_count} critical issue(s) and "
            f"{result.warning_count} warning(s) that should be addressed."
        )
    else:
        summary_text = (
            f"The scanned document scored {result.score:.1f}% and is rated "
            "Non-Compliant. "
            f"{result.critical_count} critical issue(s) require immediate attention."
        )

    bs = getattr(result, "branding_summary", {})
    vs = getattr(result, "vuln_summary", {})
    failed_validators = {f.validator for f in result.findings}
    
    def get_comp_stat_txt(name):
        return "Failed" if name in failed_validators else "Pass"

    lines = [
        "=" * 72,
        "  DOCUMENT COMPLIANCE & VALIDATION REPORT",
        "=" * 72,
        f"  Generated   : {_human_timestamp()}",
        f"  File Type   : {result.file_type.upper()}",
        f"  Page Count  : {result.page_count}",
        "=" * 72,
        "",
        "COMPLIANCE SCORE",
        "-" * 40,
        f"  Score  : {result.score:.1f}%",
        f"  Grade  : {result.grade}",
        f"  Critical Findings : {result.critical_count}",
        f"  Warnings          : {result.warning_count}",
        f"  Informational     : {result.info_count}",
        f"  Passed Checks     : {result.passed_checks}",
        "",
        "EXECUTIVE SUMMARY",
        "-" * 40,
        f"  {summary_text}",
    ]

    if bs is not None:
        lines += [
            "",
            "BRANDING VALIDATION SUMMARY",
            "-" * 40,
            f"  Primary Organization   : {bs.get('primary_org', 'Unknown')}",
            f"  Validation Status      : {bs.get('status', 'N/A')}",
            f"  Logo Present           : {'Yes' if bs.get('logo_present') else 'No'}",
            f"  Brand Consistency      : {bs.get('brand_consistency', 'N/A')}",
            f"  Total Logos Detected   : {bs.get('total_logos', 0)}",
            f"  Consistency Score      : {bs.get('consistency_score', 0):.1f}%",
            f"  Detected Organizations : {', '.join(bs.get('detected_orgs', [])) if bs.get('detected_orgs') else 'None'}",
            f"  Pages With Logos       : {bs.get('pages_containing_logos', 0)}",
        ]

    if vs:
        sev_b = vs.get("severity_breakdown", {})
        sev_str = f"Crit: {sev_b.get('Critical',0)} | Warn: {sev_b.get('Warning',0)} | Info: {sev_b.get('Information',0)}"
        lines += [
            "",
            "VULNERABILITY INTELLIGENCE SUMMARY",
            "-" * 40,
            f"  Total Vulnerabilities     : {vs.get('total', 0)}",
            f"  Matched Vulnerabilities   : {vs.get('matched', 0)}",
            f"  Unmatched Vulnerabilities : {vs.get('unmatched', 0)}",
            f"  Coverage %                : {vs.get('coverage', 0):.1f}%",
            f"  Severity Breakdown        : {sev_str}",
        ]

    lines += [
        "",
        "COMPLIANCE VALIDATION SUMMARY",
        "-" * 40,
        f"  Required Sections : {get_comp_stat_txt('Required Sections')}",
        f"  Date Validation   : {get_comp_stat_txt('Date Validation')}",
        f"  Terminology Check : {get_comp_stat_txt('Terminology Check')}",
        f"  Spelling Check    : {get_comp_stat_txt('Spelling Check')}",
        "",
        "DETAILED FINDINGS",
        "-" * 40,
    ]


    if not result.findings:
        lines.append("  No issues detected. Document appears compliant.")
    else:
        for i, f in enumerate(result.findings, 1):
            lines += [
                "",
                f"  [{i}] [{f.severity.upper()}] {f.title}",
                f"      Validator    : {f.validator}",
                f"      Location     : {f.location}",
                f"      Description  : {f.description}",
                f"      Recommend.   : {f.recommendation}",
            ]

    lines += [
        "",
        "=" * 72,
        "  END OF REPORT",
        "=" * 72,
    ]

    with open(filepath, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))

    return filepath


# ─────────────────────────────────────────────────────────────────────────────
# DOCX Export
# ─────────────────────────────────────────────────────────────────────────────
def export_docx(result: ScanResult, filename_base: str = None, filepath: str = None) -> str:
    """Generate a DOCX report using python-docx. Returns file path."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import docx.oxml

    if not filepath:
        reports_dir = get_reports_dir()
        filename_base = filename_base or "compliance_report"
        filepath = os.path.join(reports_dir, f"{filename_base}_{_timestamp()}.docx")

    doc = Document()

    # Title
    title_para = doc.add_heading("Document Compliance & Validation Report", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {_human_timestamp()}")
    doc.add_paragraph(f"File Type: {result.file_type.upper()}   |   Pages: {result.page_count}")
    doc.add_paragraph("")

    # Score section
    doc.add_heading("Compliance Score", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"Score: {result.score:.1f}%  —  {result.grade}")
    run.bold = True
    run.font.size = Pt(14)
    r, g, b = _grade_color(result.grade)
    run.font.color.rgb = RGBColor(int(r * 255), int(g * 255), int(b * 255))

    summary = doc.add_table(rows=1, cols=4)
    summary.style = "Table Grid"
    hdr = summary.rows[0].cells
    hdr[0].text = f"Critical: {result.critical_count}"
    hdr[1].text = f"Warnings: {result.warning_count}"
    hdr[2].text = f"Info: {result.info_count}"
    hdr[3].text = f"Passed: {result.passed_checks}"
    doc.add_paragraph("")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    if result.grade == "Compliant":
        summary_text = (
            "The scanned document achieved a compliance score of "
            f"{result.score:.1f}% and is rated Compliant. "
            "No critical issues were identified. Minor improvements may still apply."
        )
    elif result.grade == "Partially Compliant":
        summary_text = (
            f"The scanned document scored {result.score:.1f}% and is rated "
            "Partially Compliant. "
            f"There are {result.critical_count} critical issue(s) and "
            f"{result.warning_count} warning(s) that should be addressed."
        )
    else:
        summary_text = (
            f"The scanned document scored {result.score:.1f}% and is rated "
            "Non-Compliant. "
            f"{result.critical_count} critical issue(s) require immediate attention."
        )
    doc.add_paragraph(summary_text)
    doc.add_paragraph("")

    # Branding Validation Summary
    bs = getattr(result, "branding_summary", {})
    if bs is not None:
        doc.add_heading("Branding Validation Summary", level=1)
        b_stat = bs.get("status", "N/A")
        
        b_table = doc.add_table(rows=4, cols=4)
        b_table.style = "Table Grid"
        
        branding_rows = [
            ("Primary Organization", bs.get("primary_org", "Unknown"), "Validation Status", b_stat),
            ("Logo Present", "Yes" if bs.get("logo_present") else "No", "Brand Consistency", bs.get("brand_consistency", "N/A")),
            ("Total Logos Detected", str(bs.get("total_logos", 0)), "Consistency Score", f"{bs.get('consistency_score', 0):.1f}%"),
            ("Detected Organizations", ", ".join(bs.get("detected_orgs", [])) if bs.get("detected_orgs") else "None", "Pages With Logos", str(bs.get("pages_containing_logos", 0)))
        ]
        
        for r_idx, row_data in enumerate(branding_rows):
            row_cells = b_table.rows[r_idx].cells
            for c_idx, text in enumerate(row_data):
                p = row_cells[c_idx].paragraphs[0]
                if c_idx % 2 == 0:
                    run = p.add_run(text)
                    run.bold = True
                else:
                    run = p.add_run(text)
                    if r_idx == 0 and c_idx == 3:  # Validation Status value
                        run.bold = True
                        if text == "Pass":
                            run.font.color.rgb = RGBColor(46, 125, 50)
                        elif text == "Failed":
                            run.font.color.rgb = RGBColor(198, 40, 40)
        doc.add_paragraph("")

    # Vulnerability Intelligence Summary
    vs = getattr(result, "vuln_summary", {})
    if vs:
        doc.add_heading("Vulnerability Intelligence Summary", level=1)
        sev_b = vs.get("severity_breakdown", {})
        sev_str = f"Crit: {sev_b.get('Critical',0)} | Warn: {sev_b.get('Warning',0)} | Info: {sev_b.get('Information',0)}"
        
        v_table = doc.add_table(rows=3, cols=4)
        v_table.style = "Table Grid"
        
        vuln_rows = [
            ("Total Vulnerabilities", str(vs.get("total", 0)), "Matched Vulnerabilities", str(vs.get("matched", 0))),
            ("Unmatched Vulnerabilities", str(vs.get("unmatched", 0)), "Coverage %", f"{vs.get('coverage', 0):.1f}%"),
            ("Severity Breakdown", sev_str, "", "")
        ]
        
        for r_idx, row_data in enumerate(vuln_rows):
            row_cells = v_table.rows[r_idx].cells
            for c_idx, text in enumerate(row_data):
                p = row_cells[c_idx].paragraphs[0]
                if c_idx % 2 == 0:
                    run = p.add_run(text)
                    run.bold = True
                else:
                    run = p.add_run(text)
        doc.add_paragraph("")

    # Compliance Validation Summary
    doc.add_heading("Compliance Validation Summary", level=1)
    failed_validators = {f.validator for f in result.findings}
    
    def get_comp_stat_text(name):
        return "Failed" if name in failed_validators else "Pass"
        
    c_table = doc.add_table(rows=2, cols=4)
    c_table.style = "Table Grid"
    
    comp_rows = [
        ("Required Sections", get_comp_stat_text("Required Sections"), "Date Validation", get_comp_stat_text("Date Validation")),
        ("Terminology Check", get_comp_stat_text("Terminology Check"), "Spelling Check", get_comp_stat_text("Spelling Check"))
    ]
    
    for r_idx, row_data in enumerate(comp_rows):
        row_cells = c_table.rows[r_idx].cells
        for c_idx, text in enumerate(row_data):
            p = row_cells[c_idx].paragraphs[0]
            if c_idx % 2 == 0:
                run = p.add_run(text)
                run.bold = True
            else:
                run = p.add_run(text)
                run.bold = True
                if text == "Pass":
                    run.font.color.rgb = RGBColor(46, 125, 50)
                else:
                    run.font.color.rgb = RGBColor(198, 40, 40)
    doc.add_paragraph("")

    # Detailed Findings
    doc.add_heading("Detailed Findings", level=1)

    if not result.findings:
        doc.add_paragraph("No issues detected. Document appears compliant.")
    else:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_row = table.rows[0].cells
        for i, heading in enumerate(["Severity", "Title", "Location", "Recommendation"]):
            hdr_row[i].text = heading
            for run in hdr_row[i].paragraphs[0].runs:
                run.bold = True

        for finding in result.findings:
            row = table.add_row().cells
            row[0].text = finding.severity
            row[1].text = finding.title
            row[2].text = finding.location
            row[3].text = finding.recommendation

    intel_findings = [f for f in result.findings if getattr(f, "match_quality", "")]
    if intel_findings:
        doc.add_page_break()
        doc.add_heading("Vulnerability Match Quality & Explainability Analysis", level=1)
        
        for idx, f in enumerate(intel_findings, 1):
            doc.add_heading(f"{idx}. {f.title}", level=2)
            
            p = doc.add_paragraph()
            p.add_run("Match Quality: ").bold = True
            mq_run = p.add_run(f.match_quality)
            if f.match_quality in ("Excellent", "Strong"):
                mq_run.font.color.rgb = RGBColor(0, 128, 0)
            elif f.match_quality == "Partial":
                mq_run.font.color.rgb = RGBColor(255, 140, 0)
            else:
                mq_run.font.color.rgb = RGBColor(255, 0, 0)
                
            p.add_run("  |  Confidence Score: ").bold = True
            p.add_run(f"{getattr(f, 'confidence_score', 0)}%")
            
            if f.match_quality == "Unmatched":
                p = doc.add_paragraph()
                p.add_run("KB Match: ").bold = True
                p.add_run("No Reliable Match Found")
                
                doc.add_paragraph(f"Suggested Category: {getattr(f, 'suggested_category', 'Unknown')}")
                doc.add_paragraph("Suggested Action: Review manually and consider adding a new vulnerability definition to vulnerabilities.json.")
                
                cands = getattr(f, "top_candidates", [])
                if cands:
                    p = doc.add_paragraph()
                    p.add_run("Top Candidates:").bold = True
                    for c_idx, cand in enumerate(cands):
                        doc.add_paragraph(f"{c_idx+1}. {cand['title']} (Score: {cand['score']}/100)", style="List Bullet")
                        cp = doc.add_paragraph(f"Rejection: {cand['reason']}")
                        cp.runs[0].font.color.rgb = RGBColor(200, 0, 0)
            else:
                p = doc.add_paragraph()
                p.add_run("KB Match: ").bold = True
                p.add_run(getattr(f, 'matched_vulnerability', 'Unknown'))
                
                bd = getattr(f, "match_score_breakdown", {})
                if bd:
                    bp = doc.add_paragraph()
                    bp.add_run("Score Breakdown: ").bold = True
                    bp.add_run(f"Title ({bd.get('title',0)}/30) | Desc ({bd.get('description',0)}/30) | Rem ({bd.get('remediation',0)}/25) | Sev ({bd.get('severity',0)}/15) = {bd.get('final',0)}/100")
                
                me = getattr(f, "match_evidence", {})
                if me:
                    mp = doc.add_paragraph()
                    mp.add_run("Match Evidence:").bold = True
                    if me.get("matched_keywords"): doc.add_paragraph(f"Keywords: {', '.join(me['matched_keywords'])}", style="List Bullet")
                    if me.get("matched_description_concepts"): doc.add_paragraph(f"Description Concepts: {', '.join(me['matched_description_concepts'])}", style="List Bullet")
                    if me.get("matched_remediation_concepts"): doc.add_paragraph(f"Remediation Concepts: {', '.join(me['matched_remediation_concepts'])}", style="List Bullet")
                
                mis = getattr(f, "missing_evidence", {})
                if mis:
                    mip = doc.add_paragraph()
                    mip.add_run("Missing Evidence (Lost Points):").bold = True
                    for desc, pts in mis.items():
                        dp = doc.add_paragraph(f"[-{pts}] {desc}", style="List Bullet")
                        dp.runs[0].font.color.rgb = RGBColor(200, 0, 0)

            doc.add_paragraph("")

    doc.save(filepath)
    return filepath


# ─────────────────────────────────────────────────────────────────────────────
# PDF Export via ReportLab
# ─────────────────────────────────────────────────────────────────────────────
def export_pdf(result: ScanResult, filename_base: str = None, filepath: str = None) -> str:
    """Generate a professional PDF report using ReportLab. Returns file path."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, KeepTogether,
    )

    if not filepath:
        reports_dir = get_reports_dir()
        filename_base = filename_base or "compliance_report"
        filepath = os.path.join(reports_dir, f"{filename_base}_{_timestamp()}.pdf")

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    W, H = A4

    # Custom styles
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Heading1"],
        fontSize=22,
        textColor=colors.HexColor("#1A1C1E"),
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#5A5A5A"),
        alignment=TA_CENTER,
        spaceAfter=16,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#1E3A5F"),
        spaceBefore=14,
        spaceAfter=6,
        fontName="Helvetica-Bold",
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=9,
        leading=13,
        textColor=colors.HexColor("#2D2D2D"),
    )
    finding_desc_style = ParagraphStyle(
        "FindingDesc",
        parent=styles["Normal"],
        fontSize=8,
        leading=11,
        textColor=colors.HexColor("#3D3D3D"),
    )
    header_style = ParagraphStyle(
        "HeaderStyle",
        parent=styles["Normal"],
        fontSize=9,
        fontName="Helvetica-Bold",
        textColor=colors.white,
    )

    # Grade color
    gr, gg, gb = _grade_color(result.grade)
    grade_color = colors.Color(gr, gg, gb)

    story = []

    # ── Header ──────────────────────────────────────────────────────────────
    story.append(Paragraph("Document Compliance &amp; Validation Report", title_style))
    story.append(Paragraph(
        f"Generated: {_human_timestamp()} &nbsp;|&nbsp; "
        f"Type: {result.file_type.upper()} &nbsp;|&nbsp; Pages: {result.page_count}",
        subtitle_style,
    ))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CCCCCC")))
    story.append(Spacer(1, 0.4 * cm))

    # ── Score Card ───────────────────────────────────────────────────────────
    story.append(Paragraph("Compliance Score", section_style))

    score_data = [
        [
            Paragraph(f"<b>{result.score:.1f}%</b>", ParagraphStyle(
                "ScoreNum", parent=styles["Normal"],
                fontSize=28, textColor=grade_color, alignment=TA_CENTER,
            )),
            Paragraph(f"<b>{result.grade}</b>", ParagraphStyle(
                "Grade", parent=styles["Normal"],
                fontSize=16, textColor=grade_color, alignment=TA_CENTER,
            )),
        ]
    ]
    score_table = Table(score_data, colWidths=[8 * cm, 8 * cm])
    score_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F5F7FA")),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#F5F7FA")]),
        ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#DDDDDD")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
    ]))
    story.append(score_table)
    story.append(Spacer(1, 0.3 * cm))

    # Summary stats bar
    sev_colors = {
        "Critical":    colors.HexColor("#C62828"),
        "Warning":     colors.HexColor("#EF6C00"),
        "Information": colors.HexColor("#0288D1"),
        "Passed":      colors.HexColor("#2E7D32"),
    }

    def stat_para(label, val, col):
        return Paragraph(
            f'<font color="{col.hexval() if hasattr(col, "hexval") else "#333333"}"><b>{val}</b></font><br/>'
            f'<font size="8">{label}</font>',
            ParagraphStyle("Stat", parent=styles["Normal"], alignment=TA_CENTER, leading=14),
        )

    stat_data = [[
        stat_para("Critical", result.critical_count, sev_colors["Critical"]),
        stat_para("Warnings", result.warning_count, sev_colors["Warning"]),
        stat_para("Informational", result.info_count, sev_colors["Information"]),
        stat_para("Passed Checks", result.passed_checks, sev_colors["Passed"]),
    ]]
    stat_table = Table(stat_data, colWidths=[4 * cm] * 4)
    stat_table.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#DDDDDD")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#EEEEEE")),
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFFFFF")),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(stat_table)
    story.append(Spacer(1, 0.5 * cm))

    # ── Executive Summary ────────────────────────────────────────────────────
    story.append(Paragraph("Executive Summary", section_style))
    if result.grade == "Compliant":
        summary_text = (
            "The scanned document achieved a compliance score of "
            f"<b>{result.score:.1f}%</b> and is rated <b>Compliant</b>. "
            "No critical issues were identified. Minor improvements may still apply."
        )
    elif result.grade == "Partially Compliant":
        summary_text = (
            f"The scanned document scored <b>{result.score:.1f}%</b> and is rated "
            "<b>Partially Compliant</b>. "
            f"There are <b>{result.critical_count}</b> critical issue(s) and "
            f"<b>{result.warning_count}</b> warning(s) that should be addressed."
        )
    else:
        summary_text = (
            f"The scanned document scored <b>{result.score:.1f}%</b> and is rated "
            "<b>Non-Compliant</b>. "
            f"<b>{result.critical_count}</b> critical issue(s) require immediate attention."
        )
    story.append(Paragraph(summary_text, body_style))
    story.append(Spacer(1, 0.3 * cm))

    # ── Branding Validation Summary ──────────────────────────────────────────
    bs = getattr(result, "branding_summary", {})
    if bs is not None:
        story.append(Paragraph("Branding Validation Summary", section_style))
        b_stat = bs.get("status", "N/A")
        b_col = colors.HexColor("#2E7D32") if b_stat == "Pass" else colors.HexColor("#C62828") if b_stat == "Failed" else colors.HexColor("#1A1C1E")
        
        branding_data = [
            [Paragraph("<b>Primary Organization</b>", body_style), Paragraph(bs.get("primary_org", "Unknown"), body_style),
             Paragraph("<b>Validation Status</b>", body_style), Paragraph(f"<b><font color='{b_col.hexval()}'>{b_stat}</font></b>", body_style)],
            [Paragraph("<b>Logo Present</b>", body_style), Paragraph("Yes" if bs.get("logo_present") else "No", body_style),
             Paragraph("<b>Brand Consistency</b>", body_style), Paragraph(bs.get("brand_consistency", "N/A"), body_style)],
            [Paragraph("<b>Total Logos Detected</b>", body_style), Paragraph(str(bs.get("total_logos", 0)), body_style),
             Paragraph("<b>Consistency Score</b>", body_style), Paragraph(f"{bs.get('consistency_score', 0):.1f}%", body_style)],
            [Paragraph("<b>Detected Organizations</b>", body_style), Paragraph(", ".join(bs.get("detected_orgs", [])) if bs.get("detected_orgs") else "None", body_style),
             Paragraph("<b>Pages With Logos</b>", body_style), Paragraph(str(bs.get("pages_containing_logos", 0)), body_style)]
        ]
        
        b_table = Table(branding_data, colWidths=[4*cm, 4.5*cm, 4*cm, 4.5*cm])
        b_table.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#DDDDDD")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F5F7FA")),
            ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#F5F7FA")),
        ]))
        story.append(b_table)
        story.append(Spacer(1, 0.3 * cm))

    # ── Vulnerability Intelligence Summary ───────────────────────────────────
    vs = getattr(result, "vuln_summary", {})
    if vs:
        story.append(Paragraph("Vulnerability Intelligence Summary", section_style))
        sev_b = vs.get("severity_breakdown", {})
        sev_str = f"Crit: {sev_b.get('Critical',0)} | Warn: {sev_b.get('Warning',0)} | Info: {sev_b.get('Information',0)}"
        
        vuln_data = [
            [Paragraph("<b>Total Vulnerabilities</b>", body_style), Paragraph(str(vs.get("total", 0)), body_style),
             Paragraph("<b>Matched Vulnerabilities</b>", body_style), Paragraph(str(vs.get("matched", 0)), body_style)],
            [Paragraph("<b>Unmatched Vulnerabilities</b>", body_style), Paragraph(str(vs.get("unmatched", 0)), body_style),
             Paragraph("<b>Coverage %</b>", body_style), Paragraph(f"{vs.get('coverage', 0):.1f}%", body_style)],
            [Paragraph("<b>Severity Breakdown</b>", body_style), Paragraph(sev_str, body_style),
             Paragraph("", body_style), Paragraph("", body_style)]
        ]
        
        v_table = Table(vuln_data, colWidths=[4.5*cm, 4*cm, 4.5*cm, 4*cm])
        v_table.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#DDDDDD")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F5F7FA")),
            ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#F5F7FA")),
        ]))
        story.append(v_table)
        story.append(Spacer(1, 0.3 * cm))

    # ── Compliance Validation Summary ────────────────────────────────────────
    story.append(Paragraph("Compliance Validation Summary", section_style))
    failed_validators = {f.validator for f in result.findings}
    
    def get_comp_stat_para(name):
        if name in failed_validators: 
            return Paragraph("<b><font color='#C62828'>Failed</font></b>", body_style)
        return Paragraph("<b><font color='#2E7D32'>Pass</font></b>", body_style)

    comp_data = [
        [Paragraph("<b>Required Sections</b>", body_style), get_comp_stat_para("Required Sections"),
         Paragraph("<b>Date Validation</b>", body_style), get_comp_stat_para("Date Validation")],
        [Paragraph("<b>Terminology Check</b>", body_style), get_comp_stat_para("Terminology Check"),
         Paragraph("<b>Spelling Check</b>", body_style), get_comp_stat_para("Spelling Check")]
    ]
    
    c_table = Table(comp_data, colWidths=[4.5*cm, 4*cm, 4.5*cm, 4*cm])
    c_table.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#DDDDDD")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F5F7FA")),
        ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#F5F7FA")),
    ]))
    story.append(c_table)
    story.append(Spacer(1, 0.5 * cm))

    # ── Findings Table ───────────────────────────────────────────────────────
    story.append(Paragraph("Detailed Findings", section_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC")))
    story.append(Spacer(1, 0.2 * cm))

    if not result.findings:
        story.append(Paragraph(
            "✓ No issues were detected. The document appears fully compliant.",
            ParagraphStyle("Good", parent=styles["Normal"], fontSize=10,
                           textColor=colors.HexColor("#2E7D32")),
        ))
    else:
        header_row = [
            Paragraph("<b>#</b>", header_style),
            Paragraph("<b>Severity</b>", header_style),
            Paragraph("<b>Finding</b>", header_style),
            Paragraph("<b>Location</b>", header_style),
        ]
        table_data = [header_row]

        for i, f in enumerate(result.findings, 1):
            r2, g2, b2 = _severity_color_rgb(f.severity)
            sev_color = colors.Color(r2, g2, b2)
            sev_para = Paragraph(
                f'<font color="white"><b> {f.severity} </b></font>',
                ParagraphStyle("Sev", parent=styles["Normal"], fontSize=8,
                               backColor=sev_color, alignment=TA_CENTER),
            )
            finding_para = Paragraph(
                f"<b>{f.title}</b><br/>"
                f"<font size='7' color='#555555'>{f.description[:120]}{'...' if len(f.description) > 120 else ''}</font>",
                finding_desc_style,
            )
            loc_para = Paragraph(f.location[:60], finding_desc_style)
            table_data.append([
                Paragraph(str(i), body_style),
                sev_para,
                finding_para,
                loc_para,
            ])

        col_widths = [0.8 * cm, 2.5 * cm, 10.5 * cm, 3.5 * cm]
        findings_table = Table(table_data, colWidths=col_widths, repeatRows=1)

        row_bgs = []
        for row_idx in range(1, len(table_data)):
            bg = colors.HexColor("#FFF8F8") if row_idx % 2 == 0 else colors.white
            row_bgs.append(("BACKGROUND", (0, row_idx), (-1, row_idx), bg))

        findings_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8F9FA")]),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#DDDDDD")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            *row_bgs,
        ]))
        story.append(findings_table)

    # ── Vulnerability Intelligence Details (Phase 5E) ────────────────────────
    intel_findings = [f for f in result.findings if getattr(f, "match_quality", "")]
    if intel_findings:
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph("Vulnerability Intelligence Analysis", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC")))
        story.append(Spacer(1, 0.2 * cm))

        # Structured Finding Cards
        card_label_style = ParagraphStyle("CardLabel", parent=styles["Normal"], fontSize=9, fontName="Helvetica-Bold", textColor=colors.HexColor("#444444"))
        
        for idx, f in enumerate(intel_findings):
            card_elements = []
            
            # Extract page and paragraph from location (e.g. "Page X, Paragraph Y")
            page_val = "N/A"
            para_val = "N/A"
            if f.location and "Page" in f.location:
                parts = f.location.split(",")
                for part in parts:
                    if "Page" in part:
                        page_val = part.replace("Page", "").strip()
                    elif "Paragraph" in part:
                        para_val = part.replace("Paragraph", "").strip()

            card_data = []
            card_data.append([Paragraph("Actual Vulnerability Name:", card_label_style), Paragraph(f.title, body_style)])
            card_data.append([Paragraph("KB Match:", card_label_style), Paragraph(f"<b>{getattr(f, 'matched_vulnerability', 'Unknown')}</b>", body_style)])
            
            mq_color = "#2E7D32" if f.match_quality in ("Excellent", "Strong") else "#EF6C00" if f.match_quality == "Partial" else "#C62828"
            card_data.append([Paragraph("Confidence Score:", card_label_style), Paragraph(f"<b>{getattr(f, 'confidence_score', 0)}%</b> (<font color='{mq_color}'><b>{f.match_quality}</b></font>)", body_style)])
            
            card_data.append([Paragraph("Page Number:", card_label_style), Paragraph(page_val, body_style)])
            card_data.append([Paragraph("Paragraph Number:", card_label_style), Paragraph(para_val, body_style)])
            
            if f.match_quality == "Unmatched":
                card_data.append([Paragraph("Suggested Category:", card_label_style), Paragraph(getattr(f, 'suggested_category', 'Unknown'), body_style)])
                card_data.append([Paragraph("Suggested Action:", card_label_style), Paragraph("Review manually and consider adding a new vulnerability definition to vulnerabilities.json.", body_style)])
                
                if getattr(f, "top_candidates", []):
                    cands = []
                    for idx, cand in enumerate(f.top_candidates):
                        cands.append(f"{idx+1}. {cand['title']} (Score: {cand['score']}/100)<br/>&nbsp;&nbsp;&nbsp;<font color='#C62828'>Rejection: {cand['reason']}</font>")
                    card_data.append([Paragraph("Top Candidates:", card_label_style), Paragraph("<br/>".join(cands), finding_desc_style)])
            else:
                bd = getattr(f, "match_score_breakdown", {})
                if bd:
                    score_txt = f"Title ({bd.get('title',0)}/30) | Desc ({bd.get('description',0)}/30) | Rem ({bd.get('remediation',0)}/25) | Sev ({bd.get('severity',0)}/15) = <b>{bd.get('final',0)}/100</b>"
                    card_data.append([Paragraph("Score Breakdown:", card_label_style), Paragraph(score_txt, body_style)])
                
                me = getattr(f, "match_evidence", {})
                if me:
                    me_lines = []
                    if me.get("matched_keywords"): me_lines.append(f"• Keywords: {', '.join(me['matched_keywords'])}")
                    if me.get("matched_description_concepts"): me_lines.append(f"• Description Concepts: {', '.join(me['matched_description_concepts'])}")
                    if me.get("matched_remediation_concepts"): me_lines.append(f"• Remediation Concepts: {', '.join(me['matched_remediation_concepts'])}")
                    if me_lines:
                        card_data.append([Paragraph("Match Evidence:", card_label_style), Paragraph("<br/>".join(me_lines), finding_desc_style)])

                mis = getattr(f, "missing_evidence", {})
                if mis:
                    mis_lines = []
                    for desc, pts in mis.items():
                        mis_lines.append(f"<font color='#C62828'>• [-{pts}] {desc}</font>")
                    if mis_lines:
                        card_data.append([Paragraph("Missing Evidence:", card_label_style), Paragraph("<br/>".join(mis_lines), finding_desc_style)])
            
            card_table = Table(card_data, colWidths=[4*cm, 12*cm])
            card_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8F9FA")),
                ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#DDDDDD")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8)
            ]))
            
            card_elements.append(card_table)
            card_elements.append(Spacer(1, 0.4 * cm))
            
            if idx < len(intel_findings) - 1:
                card_elements.append(HRFlowable(width="50%", thickness=0.5, color=colors.HexColor("#E0E0E0"), spaceAfter=10))
            
            story.append(KeepTogether(card_elements))

    # ── Recommendations ──────────────────────────────────────────────────────
    if result.findings:
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph("Recommendations", section_style))
        critical = [f for f in result.findings if f.severity == "Critical"]
        warnings = [f for f in result.findings if f.severity == "Warning"]

        if critical:
            story.append(Paragraph("<b>Critical (Immediate Action Required):</b>", body_style))
            for f in critical:
                story.append(Paragraph(f"• {f.recommendation}", body_style))
        if warnings:
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("<b>Warnings (Should Be Addressed):</b>", body_style))
            for f in warnings:
                story.append(Paragraph(f"• {f.recommendation}", body_style))

    # ── Footer note ──────────────────────────────────────────────────────────
    story.append(Spacer(1, 1 * cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC")))
    story.append(Paragraph(
        "This report was generated automatically by Document Compliance &amp; Validation Checker. "
        "Results should be reviewed by a qualified professional.",
        ParagraphStyle("Footer", parent=styles["Normal"], fontSize=7.5,
                       textColor=colors.HexColor("#888888"), alignment=TA_CENTER),
    ))

    doc.build(story)
    return filepath


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────
def export_report(result: ScanResult, format_type: str, filename_base: str = "compliance_report") -> str:
    """
    Export a report to the default reports/ directory.
    format_type: "pdf" | "docx" | "txt"
    Returns the path to the generated file.
    """
    fmt = format_type.lower()
    if fmt == "pdf":
        return export_pdf(result, filename_base)
    elif fmt == "docx":
        return export_docx(result, filename_base)
    elif fmt == "txt":
        return export_txt(result, filename_base)
    else:
        raise ValueError(f"Unsupported report format: {format_type}")


def export_report_to_path(result: ScanResult, format_type: str, save_path: str) -> str:
    """
    Export a report to a specific user-chosen file path.
    Generates the content in memory and writes to the given path.
    Returns the path on success.
    """
    fmt = format_type.lower()

    if fmt == "pdf":
        return export_pdf(result, filepath=save_path)
    elif fmt == "docx":
        return export_docx(result, filepath=save_path)
    elif fmt == "txt":
        return export_txt(result, filepath=save_path)
    else:
        raise ValueError(f"Unsupported report format: {format_type}")

